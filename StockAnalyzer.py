import os
import yfinance as yf
import requests
import feedparser
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Load environment variables
load_dotenv()

# Define API key securely
API_KEY = os.getenv('GENERATIVEAI_API_KEY')

# Import necessary libraries for AI model
try:
    import google.generativeai as genai
    genai.configure(api_key=API_KEY)
except ImportError:
    print("Google Generative AI library not installed or API key not set.")
    exit()

# Define generation configurations
generation_config = {
    "temperature": 0,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

generation_config_2 = {
    "temperature": 0,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# Initialize AI models
try:
    model = genai.GenerativeModel(
        model_name="gemini-2.0-flash",
        generation_config=generation_config,
        system_instruction="You are a financial professional, not an AI chatbot.",
    )
    model_2 = genai.GenerativeModel(
        model_name="gemini-2.0-flash",
        generation_config=generation_config_2,
        system_instruction="Find the ticker of the company, if you cannot find the ticker, respond with just the letter x, also do not add the",
    )
except Exception as e:
    print(f"Failed to initialize AI models: {e}")
    exit()

def get_stock_analysis(company_name_or_ticker, term):
    dont_run = False
    fortune_500_top_100 = company_name_or_ticker.replace(" ", "+")
    query = f"{fortune_500_top_100}+stock+market"
    rss_url = f"https://news.google.com/rss/search?q={query}&hl=en-US&gl=US&ceid=US:en"
    feed = feedparser.parse(rss_url)
    news_list = []
    
    for entry in feed.entries[:50]:
        article_url = entry.link
        response = requests.get(article_url, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = soup.find_all("p")
        article_body = "\n".join([p.text for p in paragraphs[:5]])
        news_article = {
            "title": entry.title,
            "source": entry.source.title if hasattr(entry, 'source') else "Unknown",
            "url": article_url,
            "body": article_body
        }
        news_list.append(news_article)
    
    try:
        ticker = yf.Ticker(fortune_500_top_100)
        financials = ticker.financials
        if financials.empty:
            raise RuntimeError("Empty Frame")
        current_price = ticker.info['regularMarketPrice']
        historical_data = ticker.history(period="max")
        
        userinput = (
            f"Provide a detailed financial analysis for {fortune_500_top_100} at current price {str(current_price)}, "
            f"covering their financials ({financials}), historical stock prices ({historical_data}), and "
            f"include an analysis of recent news ({news_list}) into account. "
            f"Additionally, conduct research on current market trends from both macroeconomic and microeconomic perspectives. "
            f"Based on this information, assess whether it is ideal to invest in the {term} term. "
            f"Assign a letter grade (A, B, C, D, or F) reflecting the investment potential of the specified ticker. "
            f"Clearly explain your reasoning, including any risks, growth potential, and external market influences. "
            f"The final grade should be on the last line, and no additional text should follow the grade."
        )
        
        chat_session = model.start_chat(history=[])
        response = chat_session.send_message(userinput)
        response = response.text
        
        return response
    
    except:
        ticker_form = model_2.start_chat(history=[]).send_message(fortune_500_top_100)
        response = ticker_form.text.strip()
        
        if response.upper() == "X":
            print("Stock does not exist")
            dont_run = True
        else:
            ticker = yf.Ticker(str(response))
            financials = ticker.financials
            current_price = ticker.info['regularMarketPrice']
            historical_data = ticker.history(period="max")
            
            userinput = (
                f"Provide a detailed financial analysis for {fortune_500_top_100} at current price {str(current_price)}, "
                f"covering their financials ({financials}), historical stock prices ({historical_data}), and "
                f"include an analysis of recent news ({news_list}) into account. "
                f"Additionally, conduct research on current market trends from both macroeconomic and microeconomic perspectives. "
                f"Based on this information, assess whether it is ideal to invest in the {term} term. "
                f"Assign a letter grade (A, B, C, D, or F) reflecting the investment potential of the specified ticker. "
                f"Clearly explain your reasoning, including any risks, growth potential, and external market influences. "
                f"The final grade should be on the last line, and no additional text should follow the grade."
            )
            
            chat_session = model.start_chat(history=[])
            response = chat_session.send_message(userinput)
            response = response.text
            
            return response

def generate_docx_from_txt(input_filename, output_filename, image_path, image_path2, stock):
    try:
        with open(input_filename, 'r') as file:
            lines = file.readlines()
        
        doc = Document()
        general_page_paragraph = doc.add_heading(f"Stock Analysis for {stock}", level=0)
        general_page_paragraph.alignment = 1
        
        doc.add_picture(image_path, width=Inches(6))
        doc.add_paragraph("\n")
        doc.add_picture(image_path2, width=Inches(6))
        
        doc.add_page_break()
        
        table_data = []
        is_table = False
        
        for line in lines:
            stripped_line = line.strip()
            
            if not stripped_line:
                continue
            
            if stripped_line.startswith('**') and stripped_line.endswith('**'):
                if is_table:
                    add_table_to_doc(doc, table_data)
                    table_data = []
                    is_table = False
                
                header_text = stripped_line.strip('**')
                doc.add_heading(header_text, level=1)
            
            elif stripped_line.startswith('*'):
                if is_table:
                    add_table_to_doc(doc, table_data)
                    table_data = []
                    is_table = False
                
                subheader_text = stripped_line.strip('*')
                doc.add_paragraph(subheader_text)
            
            elif stripped_line.startswith('|'):
                is_table = True
                
                if '-' in stripped_line:
                    continue
                
                cells = stripped_line.strip().split('|')[1:-1]
                cells = [cell.strip() for cell in cells]
                table_data.append(cells)
            
            else:
                if is_table:
                    add_table_to_doc(doc, table_data)
                    table_data = []
                    is_table = False
                
                doc.add_paragraph(stripped_line)
        
        if is_table:
            add_table_to_doc(doc, table_data)
        
        doc.save(output_filename)
        print(f"Successfully generated {output_filename} from {input_filename}.")
    
    except FileNotFoundError:
        print(f"Error: The file {input_filename} was not found.")
    
    except Exception as e:
        print(f"An error occurred: {e}")

def add_table_to_doc(doc, table_data):
    if not table_data:
        return
    
    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
    
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell

def plot_stock_prices(ticker, stock_name):
    try:
        historical_data_ytd = ticker.history(period="ytd")
        historical_data_max = ticker.history(period="max")
        
        start_price_ytd = historical_data_ytd['Close'].iloc[0]
        end_price_ytd = historical_data_ytd['Close'].iloc[-1]
        overall_change_ytd = end_price_ytd - start_price_ytd
        
        if overall_change_ytd > 0:
            line_color_ytd = 'green'
            fill_color_ytd = 'lightgreen'
        elif overall_change_ytd < 0:
            line_color_ytd = 'red'
            fill_color_ytd = 'lightcoral'
        else:
            line_color_ytd = 'gray'
            fill_color_ytd = 'lightgray'
        
        plt.figure(figsize=(10, 5))
        plt.plot(historical_data_ytd.index, historical_data_ytd['Close'], label='Close Price', color=line_color_ytd)
        plt.fill_between(historical_data_ytd.index, historical_data_ytd['Close'], color=fill_color_ytd, alpha=0.5)
        
        min_price_ytd = historical_data_ytd['Close'].min()
        max_price_ytd = historical_data_ytd['Close'].max()
        price_range_ytd = max_price_ytd - min_price_ytd
        y_min_ytd = min_price_ytd - (price_range_ytd * 0.1)
        y_max_ytd = max_price_ytd + (price_range_ytd * 0.1)
        plt.ylim(y_min_ytd, y_max_ytd)
        
        plt.title(f"{stock_name} Closing Prices since year to date")
        plt.xlabel('Date')
        plt.ylabel('Price (USD)')
        plt.legend()
        plt.savefig("stock_price_ytd.png")
        
        start_price_max = historical_data_max['Close'].iloc[0]
        end_price_max = historical_data_max['Close'].iloc[-1]
        overall_change_max = end_price_max - start_price_max
        
        if overall_change_max > 0:
            line_color_max = 'green'
            fill_color_max = 'lightgreen'
        elif overall_change_max < 0:
            line_color_max = 'red'
            fill_color_max = 'lightcoral'
        else:
            line_color_max = 'gray'
            fill_color_max = 'lightgray'
        
        plt.figure(figsize=(10, 5))
        plt.plot(historical_data_max.index, historical_data_max['Close'], label='Close Price', color=line_color_max)
        plt.fill_between(historical_data_max.index, historical_data_max['Close'], color=fill_color_max, alpha=0.5)
        
        plt.title(f"{stock_name} Closing Prices since inception")
        plt.xlabel('Date')
        plt.ylabel('Price (USD)')
        plt.legend()
        plt.savefig("stock_price_max.png")
    
    except Exception as e:
        print(f"Failed to plot stock prices: {e}")

def main():
    stock_name = input("Enter Stock Name or Ticker: ")
    
    try:
        ticker = yf.Ticker(stock_name.upper())
        price = ticker.info.get('regularMarketPrice')
        
        if price is None:
            raise Exception("No valid stock price found")
        
    except:
        ticker_form = model_2.start_chat(history=[]).send_message(stock_name).text.strip()
        
        if ticker_form.upper() == "X":
            print("Stock does not exist")
            return
        
        ticker = yf.Ticker(ticker_form.upper())
    
    print(f"Stock Price: {ticker.info.get('regularMarketPrice')}")
    
    analysis = get_stock_analysis(stock_name, "Long")
    print(analysis)
    
    with open("stock_analysis.txt", "w") as file:
        file.write(analysis)
    
    plot_stock_prices(ticker, ticker.info.get('shortName', 'N/A'))
    
    generate_docx_from_txt("stock_analysis.txt", "stock_analysis.docx", "stock_price_max.png", "stock_price_ytd.png", ticker.info.get('shortName', 'N/A'))

if __name__ == "__main__":
    main()
